import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('Multimodal_Deepfake_Detection_Final_Report_v2.docx')

# ── Heading definitions: (exact text start, outline level 0=H1 1=H2 2=H3) ──
LEVEL0 = [  # Main sections
    'INTRODUCTION', 'OBJECTIVE OF THIS PROJECT', 'PROPOSED METHODOLOGY',
    'HARDWARE AND SOFTWARE REQUIREMENTS', 'RESULT ANALYSIS',
    'LIMITATIONS', 'FUTURE SCOPE',
]
LEVEL1 = [  # Sub-sections
    '3.1', '3.2', '4.1', '4.2',
    '5.1', '5.2', '5.3', '5.4', '5.5', '5.6', '5.7  ',
]
LEVEL2 = [  # Sub-sub-sections
    'Step 1:', 'Step 2:', 'Step 3:', 'Step 4:', 'Step 5:',
    'Step 6:', 'Step 7:', 'Step 8:', 'Step 9:', 'Step 10:',
    '5.7.1', '5.7.2', '5.7.3', '5.7.4', '5.7.5',
    '5.7.6', '5.7.7', '5.7.8', '5.7.9', '5.7.10',
]

def set_outline_level(para, level):
    """Apply outline level to paragraph so TOC can detect it."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    # Remove existing outlineLvl
    for el in pPr.findall(qn('w:outlineLvl')):
        pPr.remove(el)
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level))
    pPr.append(outlineLvl)

applied = 0
for para in doc.paragraphs:
    t = para.text.strip()
    if not t:
        continue
    if any(t == h or t.startswith(h) for h in LEVEL0):
        set_outline_level(para, 0)
        print(f'H1: {t[:60]}')
        applied += 1
    elif any(t.startswith(h) for h in LEVEL1):
        set_outline_level(para, 1)
        print(f'  H2: {t[:60]}')
        applied += 1
    elif any(t.startswith(h) for h in LEVEL2):
        set_outline_level(para, 2)
        print(f'    H3: {t[:60]}')
        applied += 1

print(f'\nOutline levels applied to {applied} paragraphs.')

# ── Also update the TOC field to use outline levels ──
# Find TOC field instrText and update it
for para in doc.paragraphs:
    for elem in para._element.iter(qn('w:instrText')):
        if 'TOC' in (elem.text or ''):
            elem.text = ' TOC \\o "1-3" \\h \\z \\u '
            print('TOC field instruction updated.')
            break

doc.save('Multimodal_Deepfake_Detection_Final_Report_v3.docx')
print('Saved successfully.')
