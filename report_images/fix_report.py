import sys, os, shutil
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

doc = Document('Multimodal_Deepfake_Detection_Final_Report_v2.docx')

# ══════════════════════════════════════════════════════════════
# TASK 1 — CREATE report_images FOLDER AND MOVE FILES
# ══════════════════════════════════════════════════════════════
print('=== TASK 1: Moving images to report_images/ ===')
os.makedirs('report_images', exist_ok=True)

move_files = (
    [f for f in os.listdir('.') if f.endswith('.png')] +
    ['architecture_diagram.xml',
     'gen_stage_images.py',
     'insert_stage_images.py',
     'fix_report.py']   # move self last
)

for fname in move_files:
    if os.path.exists(fname) and fname != 'fix_report.py':
        shutil.move(fname, os.path.join('report_images', fname))
        print(f'  Moved: {fname}')

print()

# ══════════════════════════════════════════════════════════════
# TASK 2 — FIX PAGE NUMBERING
#   Section 0 (Cover)       → no footer
#   Section 1 (TOC)         → no footer
#   Section 2 (Introduction)→ dynamic page number starting at 1
# ══════════════════════════════════════════════════════════════
print('=== TASK 2: Fixing page numbering ===')

def clear_footer(section):
    """Remove all content from a section footer."""
    footer = section.footer
    for para in footer.paragraphs:
        for run in para.runs:
            run.text = ''
        # Remove all child elements except pPr
        for elem in list(para._element):
            if elem.tag != qn('w:pPr'):
                para._element.remove(elem)
    print(f'  Footer cleared.')

def set_dynamic_page_footer(section, left_text='Savan Patel', right_text='25MDS013'):
    """Replace footer with: left_text  [PAGE]  right_text, starting at page 1."""
    footer = section.footer
    # Clear existing
    for para in footer.paragraphs:
        p_elem = para._element
        for child in list(p_elem):
            p_elem.remove(child)

    # Build: "Savan Patel  [TAB]  {PAGE}  [TAB]  25MDS013"
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Left text
    r1 = OxmlElement('w:r')
    t1 = OxmlElement('w:t')
    t1.text = left_text + '\t'
    r1.append(t1)
    para._element.append(r1)

    # PAGE field
    r2 = OxmlElement('w:r')
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r2.append(fld_begin)
    para._element.append(r2)

    r3 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' PAGE '
    r3.append(instr)
    para._element.append(r3)

    r4 = OxmlElement('w:r')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r4.append(fld_end)
    para._element.append(r4)

    # Right text
    r5 = OxmlElement('w:r')
    t5 = OxmlElement('w:t')
    t5.text = '\t' + right_text
    r5.append(t5)
    para._element.append(r5)

    print(f'  Dynamic page number footer set.')

def set_page_number_start(section, start=1):
    """Set section page number to restart at given value."""
    sectPr = section._sectPr
    # Remove existing pgNumType if present
    for el in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(el)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)
    print(f'  Page number restart set to {start}.')

sections = doc.sections
print(f'Total sections: {len(sections)}')

# Section 0 — Cover page: no footer
print('Section 0 (Cover):')
clear_footer(sections[0])

# Section 1 — TOC page: no footer
print('Section 1 (TOC):')
clear_footer(sections[1])

# Section 2 — Introduction onwards: dynamic page number starting at 1
print('Section 2 (Introduction):')
set_dynamic_page_footer(sections[2])
set_page_number_start(sections[2], start=1)

print()

# ══════════════════════════════════════════════════════════════
# TASK 3 — REPLACE MANUAL TOC WITH WORD TOC FIELD
#   Delete paragraphs 18-45 (TOC entries)
#   Insert: heading "Table of Contents" + TOC field
# ══════════════════════════════════════════════════════════════
print('=== TASK 3: Replacing manual TOC with Word TOC field ===')

# Find TOC heading paragraph
toc_start = None
toc_end   = None
intro_text = 'INTRODUCTION'

for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Table of Contents' and toc_start is None:
        toc_start = i
    # TOC ends just before the blank line before INTRODUCTION
    if toc_start and p.text.strip() == intro_text and toc_end is None:
        toc_end = i - 1
        break

print(f'TOC range: paragraphs {toc_start} to {toc_end}')

# Collect elements to delete (TOC entries only, keep heading)
paras_to_delete = doc.paragraphs[toc_start + 1 : toc_end + 1]
print(f'Deleting {len(paras_to_delete)} TOC entry paragraphs...')

for p in paras_to_delete:
    p._element.getparent().remove(p._element)

# Now insert TOC field after the "Table of Contents" heading
toc_heading = None
for p in doc.paragraphs:
    if p.text.strip() == 'Table of Contents':
        toc_heading = p
        break

# Build TOC field paragraph
toc_p = OxmlElement('w:p')

# Add paragraph properties for TOC style
pPr = OxmlElement('w:pPr')
pStyle = OxmlElement('w:pStyle')
pStyle.set(qn('w:val'), 'Normal')
pPr.append(pStyle)
toc_p.append(pPr)

# begin fldChar
r_begin = OxmlElement('w:r')
fc_begin = OxmlElement('w:fldChar')
fc_begin.set(qn('w:fldCharType'), 'begin')
fc_begin.set(qn('w:dirty'), 'true')
r_begin.append(fc_begin)
toc_p.append(r_begin)

# instrText: TOC field with headings 1-3, hyperlinks, no page numbers visible initially
r_instr = OxmlElement('w:r')
instr = OxmlElement('w:instrText')
instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
instr.text = ' TOC \\o "1-3" \\h \\z \\u '
r_instr.append(instr)
toc_p.append(r_instr)

# separate
r_sep = OxmlElement('w:r')
fc_sep = OxmlElement('w:fldChar')
fc_sep.set(qn('w:fldCharType'), 'separate')
r_sep.append(fc_sep)
toc_p.append(r_sep)

# placeholder text
r_ph = OxmlElement('w:r')
rPr = OxmlElement('w:rPr')
color = OxmlElement('w:color')
color.set(qn('w:val'), '808080')
rPr.append(color)
r_ph.append(rPr)
t_ph = OxmlElement('w:t')
t_ph.text = '[Right-click here and select Update Field to generate Table of Contents]'
r_ph.append(t_ph)
toc_p.append(r_ph)

# end fldChar
r_end = OxmlElement('w:r')
fc_end = OxmlElement('w:fldChar')
fc_end.set(qn('w:fldCharType'), 'end')
r_end.append(fc_end)
toc_p.append(r_end)

# Insert after TOC heading
toc_heading._element.addnext(toc_p)
print('TOC field inserted.')

print()

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
doc.save('Multimodal_Deepfake_Detection_Final_Report_v2.docx')
print('=== Report saved: Multimodal_Deepfake_Detection_Final_Report_v2.docx ===')

# Move self to report_images after saving
if os.path.exists('fix_report.py'):
    shutil.move('fix_report.py', 'report_images/fix_report.py')
    print('Moved fix_report.py to report_images/')
