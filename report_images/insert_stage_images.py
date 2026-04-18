import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

doc = Document('Multimodal_Deepfake_Detection_Final_Report_v2.docx')

def add_image_after(ref_para, img_path, width):
    new_p = OxmlElement('w:p')
    ref_para._element.addnext(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            new_para = p; break
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(img_path, width=Inches(width))
    return new_para

def add_caption_after(ref_para, text):
    new_p = OxmlElement('w:p')
    ref_para._element.addnext(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            new_para = p; break
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return new_para

# Map: output line keyword → (image file, caption text, width)
stage_images = [
    ('Output: Raw MP4 video file with corresponding binary label',
     'out_step1_input.png',
     'Output Image 1   Sample input video frame from LAV-DF dataset (Video: 000032.mp4)',
     4.5),
    ('Output: 16 RGB frames (visual) + 4-second WAV audio',
     'out_step2_extraction.png',
     'Output Image 2   16 extracted face keyframes (top) and audio waveform energy (bottom)',
     6.0),
    ('Output: Face tensor (16, 224, 224, 3) uint8 + Mel spectrogram',
     'out_step3_preprocess.png',
     'Output Image 3   MTCNN face crops (left) and Mel Spectrogram 128x126 (right)',
     6.0),
    ('Output: Normalized float32 tensors ready for backbone feature extraction',
     'out_step4_norm.png',
     'Output Image 4   Three-stage normalization: Raw Uint8 >> Float32 >> ImageNet Normalized',
     6.0),
    ('Output: Visual feature vector (B, D_visual) + Audio feature vector',
     'out_step5_features.png',
     'Output Image 5   Visual feature vector Fv (512-d, red) and Audio feature vector Fa (128-d, green)',
     6.0),
    ('Output: Fused feature vector (B, D_fused) combining visual and audio',
     'out_step6_fusion.png',
     'Output Image 6   Fusion output: Fv (512-d) concatenated with Fa (128-d) to form 640-dim vector',
     6.0),
    ('Output: Logits (B, 2)',
     'out_step7_classification.png',
     'Output Image 7   Softmax probability output for a REAL video (left) and FAKE video (right)',
     6.0),
    ('Output: Best model checkpoint (.pth) saved at the epoch with minimum validation loss',
     'out_step8_training.png',
     'Output Image 8   M1 training curves: Cross-Entropy Loss (left) and Accuracy (right) over 10 epochs',
     6.0),
    ('Output: Per-model and per-class evaluation metrics + confusion matrices',
     'out_step9_evaluation.png',
     'Output Image 9   M1 Confusion Matrix on dev set (5,000 samples) showing per-class predictions',
     4.5),
    ('Output: Ranked model comparison with per-pair winners',
     'out_step10_comparison.png',
     'Output Image 10   Final model ranking by Validation Accuracy (left) and F1-Score (right)',
     6.0),
]

inserted = 0
for keyword, img_file, cap_text, width in stage_images:
    found = None
    for p in doc.paragraphs:
        if keyword in p.text:
            found = p; break
    if found:
        img_para = add_image_after(found, img_file, width)
        cap_para = add_caption_after(img_para, cap_text)
        print(f'Inserted: {cap_text[:60]}...')
        inserted += 1
    else:
        print(f'WARNING: Could not find paragraph containing: {keyword[:60]}')

print(f'\nTotal inserted: {inserted}/10')
doc.save('Multimodal_Deepfake_Detection_Final_Report_v2.docx')
print('Saved: Multimodal_Deepfake_Detection_Final_Report_v2.docx')
